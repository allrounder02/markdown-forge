from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from config.schema import Config

def hex_to_rgb(hex_color: str):
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def apply_style(run, style_config):
    if not style_config:
        return
    
    if style_config.size:
        # Assume pt if not specified, strip 'pt'
        size = float(style_config.size.replace('pt', ''))
        run.font.size = Pt(size)
    
    if style_config.color:
        rgb = hex_to_rgb(style_config.color)
        run.font.color.rgb = RGBColor(*rgb)
        
    if style_config.weight == 'bold':
        run.font.bold = True

def render_docx(tokens, config: Config, output_path: str):
    doc = Document()
    
    # Set margins (basic approximation)
    section = doc.sections[0]
    # Parsing '2cm' to inches/emu is needed for precise control, 
    # skipping for brevity in this MVP, relying on defaults or simple conversion if needed.

    # Iterate tokens
    # This is a simplified traversal. A full implementation would need a recursive visitor or a stack.
    # For markdown-it-py, tokens are a flat list.
    
    i = 0
    while i < len(tokens):
        token = tokens[i]
        
        if token.type == 'heading_open':
            level = int(token.tag[1])
            # Get content
            content = ""
            i += 1
            while tokens[i].type != 'heading_close':
                if tokens[i].type == 'inline':
                    content = tokens[i].content
                i += 1
            
            # Add heading
            p = doc.add_heading(content, level=level)
            
            # Apply style from config
            style_key = f"h{level}"
            style_config = getattr(config.styles, style_key, None)
            if style_config:
                for run in p.runs:
                    apply_style(run, style_config)
                    
        elif token.type == 'paragraph_open':
            p = doc.add_paragraph()
            i += 1
            while tokens[i].type != 'paragraph_close':
                if tokens[i].type == 'inline':
                    # Process inline children (bold, italic, etc)
                    # For simplicity, just dumping text. 
                    # Real implementation needs to parse inline children.
                    p.add_run(tokens[i].content)
                i += 1
                
        elif token.type == 'fence': # Code block
            p = doc.add_paragraph()
            p.style = 'Quote' # Use Quote or a custom style for code
            run = p.add_run(token.content)
            run.font.name = config.fonts.code
        
        elif token.type == 'page_break':
            # Insert a page break
            doc.add_page_break()
            
        # ... handle lists, tables, etc.
        
        i += 1
        
    doc.save(output_path)
